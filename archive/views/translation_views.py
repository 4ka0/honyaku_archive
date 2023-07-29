from django.contrib import messages
from django.contrib.auth.mixins import LoginRequiredMixin
from django.http import HttpResponseRedirect
from django.shortcuts import render
from django.urls import reverse_lazy
from django.views.generic import UpdateView, View

from docx import Document  # For reading docx files
from translate.storage.tmx import tmxfile  # For reading tmx files (from translate-toolkit)

from ..forms.translation_forms import TranslationUpdateForm, TranslationUploadForm
from ..models import Item, Resource


class TranslationUpdateView(LoginRequiredMixin, UpdateView):
    model = Resource
    template_name = "translation_update.html"
    form_class = TranslationUpdateForm

    def form_valid(self, form):
        updated_translation = form.save(commit=False)
        updated_translation.updated_by = self.request.user
        updated_translation.save()
        return HttpResponseRedirect(updated_translation.get_absolute_url())


class TranslationUploadView(LoginRequiredMixin, View):
    form_class = TranslationUploadForm
    template_name = "translation_upload.html"

    def get(self, request, *args, **kwargs):
        form = self.form_class()
        return render(request, self.template_name, {"form": form})

    def post(self, request, *args, **kwargs):
        if "cancel" in request.POST:
            if request.GET.get("previous_url"):
                previous_url = request.GET.get("previous_url")
                return HttpResponseRedirect(previous_url)

        form = TranslationUploadForm(request.POST, request.FILES)
        if form.is_valid():
            resource_obj = Resource(
                upload_file=form.cleaned_data["upload_file"],
                title=form.cleaned_data["title"],
                translator=form.cleaned_data["translator"],
                field=form.cleaned_data["field"],
                client=form.cleaned_data["client"],
                notes=form.cleaned_data["notes"],
                created_by=request.user,
                resource_type="TRANSLATION",
            )
            successful = build_items(request, resource_obj)
            if successful:
                return HttpResponseRedirect(resource_obj.get_absolute_url())
            return HttpResponseRedirect(reverse_lazy("home"))

        return render(request, self.template_name, {"form": form})


def build_items(request, resource_obj):
    """
    Helper method for TranslationUploadView.
    Chooses appropriate parser to parse the uploaded translation file, and
    builds Item objects from the parsed content.
    """

    import time

    st = time.time()  # Test-related

    # Select appropriate parser
    if resource_obj.upload_file.path.endswith(".tmx"):
        new_items = tmx_parser(resource_obj)
    else:
        new_items = docx_parser(request, resource_obj)

    # Test-related
    et = time.time()
    elapsed_time = et - st
    print("**********")
    print(elapsed_time)
    print("**********")

    # Save content to database
    if new_items:
        resource_obj.save()
        Item.objects.bulk_create(new_items)
        resource_obj.upload_file.delete()  # Uploaded file no longer needed

        """
        # Output success message if Translation and Segment objects successfully created in database
        if Translation.objects.filter(job_number__iexact=translation_obj.job_number).exists():
            if Segment.objects.filter(
                    translation__job_number__iexact=translation_obj.job_number
                ).exists:
                messages.success(request, '翻訳のアップロードが成功しました。')
        else:
            messages.error(request, '翻訳のアップロードに失敗しました。')
        """

        return True

    else:
        return False


def tmx_parser(resource_obj):
    new_items = []
    tmx_file = tmxfile(resource_obj.upload_file)

    for node in tmx_file.unit_iter():
        # Prevent None from being entered in DB for source (TextField)
        if node.source is None:
            source_text = ""
        else:
            source_text = node.source

        # Prevent None from being entered in DB for target (TextField)
        if node.target is None:
            target_text = ""
        else:
            # translate-toolkit is used to parse the TMX file. translate-toolkit assigns a control
            # character (substitute character, code point 32) when there is an empty translation
            # segment. One approach is to handle only this character, and another approach is to
            # handle all 32 control characters included at the start of the Unicode chart. The
            # latter approach has been used as it should be more robust in theory. Control
            # characters are replaced with "", which can be handled easily/cleanly in the template.
            if (
                node.target
                and len(node.target) == 1
                and ord(node.target) in range(0, 33)
            ):
                target_text = ""
            else:
                target_text = node.target

        new_segment = Item(
            resource=resource_obj,
            source=source_text,
            target=target_text,
        )
        new_items.append(new_segment)

    return new_items


def docx_parser(request, resource_obj):
    """
    Presumes that there is one table in the uploaded file, the table contains
    two columns, the first column is the source text, and the second column is
    the target text.
    """

    document = Document(resource_obj.upload_file)
    new_items = []

    if document.tables:
        if len(document.tables) > 1:
            messages.warning(
                request, ("選択したファイルに複数のテーブルが見つかりました。\n" "最初のテーブルのみが読み込まれています。")
            )

        table = document.tables[0]

        """
        # Previous method for building segments from the table content.
        # (Too slow)

        for row in table.rows:
            if len(row.cells) >= 2:
                new_segment = Segment(
                    translation=translation_obj,
                    source=row.cells[0].text,
                    target=row.cells[1].text,
                )
                new_segments.append(new_segment)
        """

        # Extract text data from table.
        # Creates one large list of 2-string lists.
        table_data = [[cell.text for cell in row.cells] for row in table.rows]

        # Create list of Segment objects from text data.
        new_items = [
            Item(resource=resource_obj, source=row[0], target=row[1])
            for row in table_data
        ]

    else:
        messages.error(request, ("翻訳のアップロードに失敗しました。\n" "選択したファイルにテーブルが見つかりません。"))

    return new_items
