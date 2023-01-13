from django.views.generic import View, DetailView, UpdateView, DeleteView
from django.urls import reverse, reverse_lazy
from django.http import HttpResponseRedirect
from django.shortcuts import render
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib import messages
# from django.http import FileResponse

from ..models import Segment, Translation
from ..forms.translation_forms import TranslationUploadForm, TranslationUpdateForm

from translate.storage.tmx import tmxfile  # For reading tmx files (from translate-toolkit)
from docx import Document  # For reading docx files


class TranslationDetailView(LoginRequiredMixin, DetailView):
    model = Translation
    template_name = "translation_detail.html"

    def get_context_data(self, **kwargs):
        context = super(TranslationDetailView, self).get_context_data(**kwargs)
        num_of_segments = context["translation"].segments.all().count()
        context.update({
            "num_of_segments": num_of_segments,
        })
        return context


class TranslationUpdateView(LoginRequiredMixin, UpdateView):
    model = Translation
    template_name = "translation_update.html"
    form_class = TranslationUpdateForm

    def form_valid(self, form):
        obj = form.save(commit=False)
        obj.updated_by = self.request.user
        obj.save()
        return HttpResponseRedirect(obj.get_absolute_url())


class TranslationDeleteView(LoginRequiredMixin, DeleteView):
    model = Translation
    template_name = "translation_delete.html"
    success_url = reverse_lazy("home")


class TranslationUploadView(LoginRequiredMixin, View):
    form_class = TranslationUploadForm
    template_name = "translation_upload.html"

    def get(self, request, *args, **kwargs):
        form = self.form_class()
        return render(request, self.template_name, {"form": form})

    def post(self, request, *args, **kwargs):
        if "cancel" in request.POST:
            if request.GET.get("previous_url"):
                previous_url = request.GET.get("previous_url")
                return HttpResponseRedirect(previous_url)

        form = TranslationUploadForm(request.POST, request.FILES)
        if form.is_valid():
            translation_obj = Translation(
                translation_file=form.cleaned_data["translation_file"],
                job_number=form.cleaned_data["job_number"],
                translator=form.cleaned_data["translator"],
                field=form.cleaned_data["field"],
                client=form.cleaned_data["client"],
                notes=form.cleaned_data["notes"],
                created_by=request.user,
                type="translation",
            )
            successful = build_segments(request, translation_obj)
            if successful:
                return HttpResponseRedirect(translation_obj.get_absolute_url())
            return HttpResponseRedirect(reverse("home"))

        return render(request, self.template_name, {"form": form})


def build_segments(request, translation_obj):
    """ Helper method for TranslationUploadView.
        Chooses appropriate parser to parse the uploaded translation file, and
        builds Segment objects from the parsed content. """

    # Select appropriate parser
    if translation_obj.translation_file.path.endswith(".tmx"):
        new_segments = tmx_parser(translation_obj)
    else:
        new_segments = docx_parser(request, translation_obj)

    # Save content to database
    if new_segments:
        translation_obj.save()
        Segment.objects.bulk_create(new_segments)
        translation_obj.translation_file.delete()  # Uploaded file no longer needed

        """
        # Output success message if Translation and Segment objects successfully created in database
        if Translation.objects.filter(job_number__iexact=translation_obj.job_number).exists():
            if Segment.objects.filter(translation__job_number__iexact=translation_obj.job_number).exists:
                messages.success(request, '翻訳のアップロードが成功しました。')
        else:
            messages.error(request, '翻訳のアップロードに失敗しました。')
        """

        return True

    else:
        return False


def tmx_parser(translation_obj):
    new_segments = []
    tmx_file = tmxfile(translation_obj.translation_file)

    for node in tmx_file.unit_iter():

        # translate-toolkit is used to parse the TMX file.
        # http://docs.translatehouse.org/projects/translate-toolkit/en/latest/api/storage.html#module-translate.storage.tmx
        #
        # translate-toolkit assigns a control character (substitute character,
        # code point 32) when there is an empty translation segment.
        # One approach is to handle only this character, and another approach is
        # to handle all 32 control characters included at the start of the
        # Unicode chart. The latter approach has been used as it should be more
        # robust in theory. Control characters are replaced with "", which can
        # be handled easily/cleanly in the template.

        if (node.target and len(node.target) == 1 and ord(node.target) in range(0, 33)):
            target_text = ""
        else:
            target_text = node.target

        new_segment = Segment(
            translation=translation_obj,
            source=node.source,
            target=target_text,
        )
        new_segments.append(new_segment)

    return new_segments


def docx_parser(request, translation_obj):
    """ Presumes that there is one table in the uploaded file, the table
        contains two columns, the first column is the source text, and the
        second column is the target text. """

    document = Document(translation_obj.translation_file)
    new_segments = []

    if document.tables:

        if len(document.tables) > 1:
            messages.warning(
                request,
                ('選択したファイルに複数のテーブルが見つかりました。\n'
                 '最初のテーブルのみが読み込まれています。')
            )

        table = document.tables[0]

        for row in table.rows:
            if len(row.cells) >= 2:
                new_segment = Segment(
                    translation=translation_obj,
                    source=row.cells[0].text,
                    target=row.cells[1].text,
                )
                new_segments.append(new_segment)

    else:
        messages.error(
            request,
            ('翻訳のアップロードに失敗しました。\n'
             '選択したファイルにテーブルが見つかりません。')
        )

    return new_segments
